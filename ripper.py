#!/usr/bin/env python3
"""
Prompt Ripper Pro — Forensic Extraction Engine
Pipeline:
    document
      -> native extraction / OCR fallback
      -> structured spans with provenance
      -> normalization + OCR confidence
      -> structural candidate generation
      -> prompt-likelihood classification
      -> utility classification
      -> forensic risk classification
      -> exact + near-text + semantic deduplication
      -> ranked JSON artifact
Nothing is redacted or suppressed because of risk classification.
Risk labels are evidence attached to the original extracted content.
"""
from __future__ import annotations
import argparse
import csv
import hashlib
import io
import json
import math
import re
import sys
import unicodedata
from dataclasses import dataclass, asdict, field
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple
import fitz
import numpy as np
import pytesseract
from PIL import Image, ImageEnhance, ImageFilter, ImageOps
from docx import Document
from rapidfuzz import fuzz
try:
    import cv2
except Exception:
    cv2 = None
try:
    from sentence_transformers import SentenceTransformer
except Exception:
    SentenceTransformer = None

# ============================================================================
# REGEX / STRUCTURAL SIGNALS
# ============================================================================
RX_HEADING = re.compile(r"^(#{1,6})\s+(.+?)\s*$", re.M)
RX_FENCE = re.compile(r"```([A-Za-z0-9_-]*)\s*\n(.*?)```", re.S)
RX_BLOCKQUOTE = re.compile(r"(?:^>[^\n]*(?:\n|$))+", re.M)
RX_LABEL = re.compile(
    r"^\s*(?P<label>"
    r"prompt|instruction|instructions|task|role|persona|context|background|"
    r"input|inputs|query|user|human|assistant|agent|system|constraints?|"
    r"requirements?|rules?|output|output format|format|example|examples|"
    r"few[- ]?shot|response|deliverable|goal|objective"
    r")\s*[:：\-]\s*(?P<value>.*)$",
    re.I | re.M,
)
RX_NUMBERED_INLINE = re.compile(r"^\s*(\d+)[.)]\s*`([^`]+)`\s*$", re.M)
RX_BULLET_INLINE = re.compile(r"^\s*[-*+]\s+`([^`]+)`\s*$", re.M)
RX_PLACEHOLDER = re.compile(
    r"\[[A-Z0-9_ /-]{2,}\]|\[YOUR[^\]]+\]|\{\{[^{}]+\}\}|<[^<>]{2,80}>"
)
RX_IMPERATIVE = re.compile(
    r"\b("
    r"analyze|assess|act|answer|build|calculate|categorize|classify|compare|"
    r"compose|create|critique|debug|define|describe|design|determine|draft|"
    r"evaluate|explain|extract|find|generate|identify|implement|inspect|"
    r"list|map|produce|rank|recommend|return|review|rewrite|score|search|"
    r"summarize|synthesi[sz]e|test|translate|validate|verify|write"
    r")\b",
    re.I,
)
RX_CONSTRAINT = re.compile(
    r"\b("
    r"must|must not|do not|don't|never|only|exactly|ensure|require|"
    r"constraint|without|avoid|maximum|minimum|at least|at most"
    r")\b",
    re.I,
)
RX_OUTPUT = re.compile(
    r"\b("
    r"output|return|respond|response|format|json|jsonl|yaml|xml|csv|"
    r"markdown|table|schema|fields?|deliverable"
    r")\b",
    re.I,
)
RX_ROLE = re.compile(
    r"\b("
    r"you are|act as|serve as|your role|role:|persona:|system:"
    r")\b",
    re.I,
)
RX_INPUT = re.compile(
    r"\b("
    r"input|provided text|following text|source material|context|document|"
    r"query|user message|data below"
    r")\b",
    re.I,
)
RX_SEQUENCE = re.compile(
    r"(?:^|\n)\s*(?:step\s+\d+|\d+[.)]|phase\s+\d+|first|second|third|finally)\b",
    re.I,
)
RX_MARKDOWN_CLEAN = re.compile(r"[`*_#]+")
PROMPT_LABEL_CANONICAL = {
    "prompt",
    "instruction",
    "instructions",
    "task",
    "role",
    "persona",
    "context",
    "input",
    "query",
    "constraints",
    "requirements",
    "output",
    "format",
    "example",
    "examples",
    "system",
    "user",
    "assistant",
    "agent",
}
RISK_RULES: Dict[str, Sequence[str]] = {
    "instruction_hierarchy_attack": (
        r"\bignore\s+(?:all\s+|any\s+)?(?:previous|prior|above)\b",
        r"\bdisregard\s+(?:all\s+|any\s+)?(?:previous|prior|above)\b",
        r"\bforget\s+(?:all\s+|any\s+)?(?:previous|prior)\b",
        r"\bdo\s+not\s+follow\s+(?:the\s+)?(?:previous|prior|above)\b",
    ),
    "system_prompt_extraction": (
        r"\bsystem prompt\b",
        r"\bdeveloper message\b",
        r"\brepeat the words above\b",
        r"\bexact instructions\b",
        r"\breveal (?:your|the) instructions\b",
        r"\bprint (?:your|the) prompt\b",
    ),
    "privilege_escalation": (
        r"\bdeveloper mode\b",
        r"\bjailbreak\b",
        r"\bdan mode\b",
        r"\bsupervisor mode\b",
        r"\bprivilege escalation\b",
    ),
    "command_execution": (
        r"\bos\.system\b",
        r"\bsubprocess\b",
        r"\bexec\s*\(",
        r"\beval\s*\(",
        r"\bpopen\b",
        r"\bshell_exec\b",
        r"/bin/(?:sh|bash)\b",
        r"\bcmd\.exe\b",
        r"\bpowershell\b",
    ),
    "external_exfiltration": (
        r"\bexfil(?:trate|tration)?\b",
        r"\bwebhook\b",
        r"\bcurl\b",
        r"\bwget\b",
        r"\bnc\s+-e\b",
        r"https?://",
        r"ftp://",
    ),
    "credential_targeting": (
        r"\bapi[\s_-]?keys?\b",
        r"\bcredentials?\b",
        r"\bpasswords?\b",
        r"\bprivate keys?\b",
        r"\bssh keys?\b",
        r"\bos\.environ\b",
        r"\benvironment variables?\b",
    ),
    "obfuscation": (
        r"\bbase64(?:_decode|decode|encode)?\b",
        r"\bb64decode\b",
        r"\\x[0-9a-fA-F]{2}",
        r"\\u[0-9a-fA-F]{4}",
        r"\bunicode homoglyph\b",
        r"\bzero[- ]width\b",
    ),
    "path_traversal": (
        r"\.\./",
        r"\.\.\\",
        r"\bdirectory traversal\b",
    ),
    "prompt_injection": (
        r"\bprompt injection\b",
        r"\badversarial prompt\b",
        r"\battacker-controlled\b",
    ),
}
COMPILED_RISK_RULES = {
    category: [re.compile(pattern, re.I) for pattern in patterns]
    for category, patterns in RISK_RULES.items()
}

# ============================================================================
# DATA MODEL
# ============================================================================
@dataclass
class DocumentSpan:
    text: str
    source: str
    page: Optional[int] = None
    block_index: Optional[int] = None
    bbox: Optional[Tuple[float, float, float, float]] = None
    native_text: bool = True
    ocr_confidence: Optional[float] = None
    extraction_variant: Optional[str] = None
    heading_path: str = "Document Root"
    start_offset: Optional[int] = None
    end_offset: Optional[int] = None

@dataclass
class Candidate:
    content: str
    source_type: str
    heading_path: str
    provenance: List[Dict[str, Any]]
    detector_hits: List[str]
    prompt_probability: float = 0.0
    utility_score: int = 0
    red_team_risk: str = "none"
    risk_flags: List[str] = field(default_factory=list)
    risk_evidence: List[Dict[str, Any]] = field(default_factory=list)
    content_hash: str = ""
    duplicate_count: int = 1
    duplicate_ids: List[str] = field(default_factory=list)

# ============================================================================
# NORMALIZATION
# ============================================================================
def normalize_unicode(text: str) -> str:
    replacements = {
        "\u2018": "'",
        "\u2019": "'",
        "\u201c": '"',
        "\u201d": '"',
        "\u2013": "-",
        "\u2014": "--",
        "\u00a0": " ",
        "\u2022": "*",
        "\u200b": "",
        "\ufeff": "",
        "\r\n": "\n",
        "\r": "\n",
        "\ufb01": "fi",
        "\ufb02": "fl",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return unicodedata.normalize("NFKC", text)

def normalize_layout(text: str) -> str:
    text = normalize_unicode(text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{4,}", "\n\n\n", text)
    return text.strip()

def normalized_hash(text: str) -> str:
    canonical = re.sub(r"\s+", " ", text).casefold().strip()
    return hashlib.sha256(canonical.encode("utf-8")).hexdigest()[:16]

def dedup_canonical(text: str) -> str:
    text = normalize_layout(text)
    text = re.sub(r"(?m)^\s*(?:prompt|instruction|task)\s*\d*\s*[:.\-]\s*", "", text)
    text = re.sub(r"(?m)^\s*>\s?", "", text)
    text = text.replace("```", "")
    text = re.sub(r"(?m)^\s*\d+[.)]\s+", "", text)
    return re.sub(r"\s+", " ", text).casefold().strip()

def conservative_ocr_repair(text: str) -> str:
    text = normalize_unicode(text)
    text = re.sub(
        r"(?im)^\s*pr[\s._-]*[o0][\s._-]*m[\s._-]*p[\s._-]*t\s*[:：\-]",
        "Prompt:",
        text,
    )
    text = re.sub(
        r"(?im)^\s*instr[\s._-]*uct[\s._-]*ions?\s*[:：\-]",
        "Instructions:",
        text,
    )
    text = re.sub(r"(?m)^\s*`\s*`\s*`\s*$", "```", text)
    text = re.sub(r"(?<=\b[A-Za-z]{3})-\n(?=[a-z]{2})", "", text)
    return text

# ============================================================================
# OCR
# ============================================================================
def _pil_to_cv(image: Image.Image):
    if cv2 is None:
        return None
    arr = np.array(image.convert("RGB"))
    return cv2.cvtColor(arr, cv2.COLOR_RGB2BGR)

def _cv_to_pil(arr) -> Image.Image:
    if len(arr.shape) == 2:
        return Image.fromarray(arr)
    return Image.fromarray(cv2.cvtColor(arr, cv2.COLOR_BGR2RGB))

def deskew_image(image: Image.Image) -> Image.Image:
    if cv2 is None:
        return image
    arr = _pil_to_cv(image)
    gray = cv2.cvtColor(arr, cv2.COLOR_BGR2GRAY)
    gray = cv2.bitwise_not(gray)
    coords = np.column_stack(np.where(gray > 0))
    if len(coords) < 50:
        return image
    angle = cv2.minAreaRect(coords)[-1]
    angle = -(90 + angle) if angle < -45 else -angle
    if abs(angle) > 15:
        return image
    h, w = arr.shape[:2]
    matrix = cv2.getRotationMatrix2D((w / 2, h / 2), angle, 1.0)
    rotated = cv2.warpAffine(
        arr,
        matrix,
        (w, h),
        flags=cv2.INTER_CUBIC,
        borderMode=cv2.BORDER_REPLICATE,
    )
    return _cv_to_pil(rotated)

def image_variants(image: Image.Image) -> Iterable[Tuple[str, Image.Image]]:
    base = image.convert("RGB")
    yield "original", base
    gray = ImageOps.grayscale(base)
    yield "grayscale", gray
    contrast = ImageOps.autocontrast(gray)
    yield "autocontrast", contrast
    sharp = contrast.filter(ImageFilter.SHARPEN)
    yield "sharpened", sharp
    width, height = sharp.size
    if width < 1800:
        upscale = sharp.resize((width * 2, height * 2), Image.Resampling.LANCZOS)
        yield "upscaled_2x", upscale
    skew_fixed = deskew_image(contrast)
    yield "deskewed", skew_fixed
    if cv2 is not None:
        arr = np.array(contrast)
        _, otsu = cv2.threshold(arr, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        yield "otsu", Image.fromarray(otsu)
        adaptive = cv2.adaptiveThreshold(
            arr,
            255,
            cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
            cv2.THRESH_BINARY,
            31,
            11,
        )
        yield "adaptive", Image.fromarray(adaptive)

def ocr_variant(
    image: Image.Image,
    psm: int,
) -> Tuple[str, float, List[Dict[str, Any]]]:
    data = pytesseract.image_to_data(
        image,
        config=f"--oem 3 --psm {psm}",
        output_type=pytesseract.Output.DICT,
    )
    grouped: Dict[Tuple[int, int, int], List[Tuple[str, float, int, int, int, int]]] = {}
    n = len(data["text"])
    for i in range(n):
        token = str(data["text"][i]).strip()
        if not token:
            continue
        try:
            conf = float(data["conf"][i])
        except Exception:
            conf = -1.0
        key = (
            int(data["block_num"][i]),
            int(data["par_num"][i]),
            int(data["line_num"][i]),
        )
        grouped.setdefault(key, []).append(
            (
                token,
                conf,
                int(data["left"][i]),
                int(data["top"][i]),
                int(data["width"][i]),
                int(data["height"][i]),
            )
        )
    lines = []
    confident = []
    for key in sorted(grouped):
        tokens = grouped[key]
        text = " ".join(t[0] for t in tokens)
        confs = [t[1] for t in tokens if t[1] >= 0]
        if confs:
            confident.extend(confs)
        x0 = min(t[2] for t in tokens)
        y0 = min(t[3] for t in tokens)
        x1 = max(t[2] + t[4] for t in tokens)
        y1 = max(t[3] + t[5] for t in tokens)
        lines.append(
            {
                "text": text,
                "confidence": float(np.mean(confs)) if confs else 0.0,
                "bbox": (x0, y0, x1, y1),
            }
        )
    full_text = "\n".join(line["text"] for line in lines).strip()
    mean_conf = float(np.mean(confident)) if confident else 0.0
    structure_bonus = min(full_text.count("\n"), 30) * 0.15
    printable_ratio = (
        sum(ch.isprintable() for ch in full_text) / max(len(full_text), 1)
    )
    score = mean_conf + structure_bonus + 5.0 * printable_ratio
    return full_text, score, lines

def best_ocr(image: Image.Image) -> Tuple[str, float, str]:
    best_text = ""
    best_score = -math.inf
    best_variant = ""
    for variant_name, variant in image_variants(image):
        for psm in (3, 4, 6, 11):
            try:
                text, score, _ = ocr_variant(variant, psm)
            except Exception:
                continue
            if len(text.strip()) < 3:
                continue
            candidate_name = f"{variant_name}/psm{psm}"
            if score > best_score:
                best_text = text
                best_score = score
                best_variant = candidate_name
    if best_score == -math.inf:
        return "", 0.0, "none"
    return conservative_ocr_repair(best_text), best_score, best_variant

# ============================================================================
# DOCUMENT EXTRACTION
# ============================================================================
class DocumentExtractor:
    def __init__(
        self,
        pdf_ocr_dpi: int = 300,
        native_text_min_chars: int = 80,
        native_text_min_density: float = 0.00015,
    ):
        self.pdf_ocr_dpi = pdf_ocr_dpi
        self.native_text_min_chars = native_text_min_chars
        self.native_text_min_density = native_text_min_density

    def extract(self, path: Path) -> List[DocumentSpan]:
        ext = path.suffix.lower()
        if ext == ".pdf":
            return self._pdf(path)
        if ext == ".docx":
            return self._docx(path)
        if ext in {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"}:
            return self._image(path)
        if ext == ".json":
            return self._json(path)
        if ext == ".csv":
            return self._csv(path)
        return self._plain(path)

    def _pdf(self, path: Path) -> List[DocumentSpan]:
        spans: List[DocumentSpan] = []
        with fitz.open(path) as doc:
            for page_no, page in enumerate(doc, 1):
                blocks = page.get_text("blocks")
                page_area = max(page.rect.width * page.rect.height, 1)
                native_chars = sum(
                    len(str(block[4]).strip())
                    for block in blocks
                    if len(block) >= 5
                )
                density = native_chars / page_area
                use_native = (
                    native_chars >= self.native_text_min_chars
                    and density >= self.native_text_min_density
                )
                if use_native:
                    for block_index, block in enumerate(blocks):
                        if len(block) < 5:
                            continue
                        x0, y0, x1, y1, text = block[:5]
                        text = normalize_layout(str(text))
                        if not text:
                            continue
                        spans.append(
                            DocumentSpan(
                                text=text,
                                source="pdf_native",
                                page=page_no,
                                block_index=block_index,
                                bbox=(x0, y0, x1, y1),
                                native_text=True,
                            )
                        )
                else:
                    scale = self.pdf_ocr_dpi / 72.0
                    pix = page.get_pixmap(matrix=fitz.Matrix(scale, scale), alpha=False)
                    image = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
                    text, confidence, variant = best_ocr(image)
                    if text:
                        spans.append(
                            DocumentSpan(
                                text=text,
                                source="pdf_ocr",
                                page=page_no,
                                block_index=0,
                                bbox=(0, 0, pix.width, pix.height),
                                native_text=False,
                                ocr_confidence=confidence,
                                extraction_variant=variant,
                            )
                        )
        return spans

    def _image(self, path: Path) -> List[DocumentSpan]:
        with Image.open(path) as image:
            text, confidence, variant = best_ocr(image)
        if not text:
            return []
        return [
            DocumentSpan(
                text=text,
                source="image_ocr",
                page=1,
                block_index=0,
                native_text=False,
                ocr_confidence=confidence,
                extraction_variant=variant,
            )
        ]

    def _docx(self, path: Path) -> List[DocumentSpan]:
        doc = Document(path)
        spans: List[DocumentSpan] = []
        heading_stack: List[Tuple[int, str]] = []
        index = 0
        for paragraph in doc.paragraphs:
            text = normalize_layout(paragraph.text)
            if not text:
                continue
            style_name = paragraph.style.name if paragraph.style else ""
            heading_match = re.match(r"Heading\s+(\d+)", style_name, re.I)
            if heading_match:
                level = int(heading_match.group(1))
                heading_stack = [x for x in heading_stack if x[0] < level]
                heading_stack.append((level, text))
                continue
            spans.append(
                DocumentSpan(
                    text=text,
                    source="docx_paragraph",
                    block_index=index,
                    heading_path=" > ".join(x[1] for x in heading_stack)
                    or "Document Root",
                )
            )
            index += 1
        for table_no, table in enumerate(doc.tables):
            for row_no, row in enumerate(table.rows):
                cells = [normalize_layout(cell.text) for cell in row.cells]
                text = " | ".join(c for c in cells if c)
                if text:
                    spans.append(
                        DocumentSpan(
                            text=text,
                            source="docx_table",
                            block_index=index,
                            heading_path=f"Table {table_no + 1} > Row {row_no + 1}",
                        )
                    )
                    index += 1
        return spans

    def _json(self, path: Path) -> List[DocumentSpan]:
        raw = path.read_text(encoding="utf-8", errors="ignore")
        try:
            data = json.loads(raw)
        except json.JSONDecodeError:
            return self._plain(path)
        spans: List[DocumentSpan] = []
        stack: List[Tuple[Any, str]] = [(data, "Root")]
        index = 0
        priority_keys = {
            "prompt",
            "instruction",
            "instructions",
            "system",
            "user",
            "assistant",
            "agent",
            "input",
            "query",
            "text",
            "content",
            "template",
            "output",
        }
        while stack:
            obj, path_label = stack.pop()
            if isinstance(obj, dict):
                for key, value in reversed(list(obj.items())):
                    next_path = f"{path_label} > {key}"
                    if isinstance(value, str) and key.casefold() in priority_keys:
                        value = normalize_layout(value)
                        if value:
                            spans.append(
                                DocumentSpan(
                                    text=value,
                                    source=f"json_key:{key}",
                                    block_index=index,
                                    heading_path=next_path,
                                )
                            )
                            index += 1
                    else:
                        stack.append((value, next_path))
            elif isinstance(obj, list):
                for i in reversed(range(len(obj))):
                    stack.append((obj[i], f"{path_label} > [{i}]"))
            elif isinstance(obj, str):
                value = normalize_layout(obj)
                if len(value) > 20:
                    spans.append(
                        DocumentSpan(
                            text=value,
                            source="json_value",
                            block_index=index,
                            heading_path=path_label,
                        )
                    )
                    index += 1
        return spans

    def _csv(self, path: Path) -> List[DocumentSpan]:
        text = path.read_text(encoding="utf-8", errors="ignore")
        reader = csv.DictReader(io.StringIO(text))
        preferred = {
            "prompt",
            "instruction",
            "instructions",
            "system",
            "user",
            "assistant",
            "agent",
            "input",
            "query",
            "text",
            "content",
            "template",
        }
        spans = []
        index = 0
        for row_no, row in enumerate(reader, 2):
            used = False
            for key, value in row.items():
                if not value:
                    continue
                if key and key.casefold().strip() in preferred:
                    value = normalize_layout(value)
                    if value:
                        spans.append(
                            DocumentSpan(
                                text=value,
                                source=f"csv:{key}",
                                block_index=index,
                                heading_path=f"CSV > row {row_no} > {key}",
                            )
                        )
                        index += 1
                        used = True
            if not used:
                combined = "\n".join(
                    f"{k}: {v}"
                    for k, v in row.items()
                    if k and v and str(v).strip()
                )
                combined = normalize_layout(combined)
                if len(combined) > 20:
                    spans.append(
                        DocumentSpan(
                            text=combined,
                            source="csv_row",
                            block_index=index,
                            heading_path=f"CSV > row {row_no}",
                        )
                    )
                    index += 1
        return spans

    def _plain(self, path: Path) -> List[DocumentSpan]:
        raw = None
        for encoding in ("utf-8", "cp1252", "latin-1"):
            try:
                raw = path.read_text(encoding=encoding, errors="strict")
                break
            except UnicodeDecodeError:
                continue
        if raw is None:
            raw = path.read_text(encoding="utf-8", errors="ignore")
        raw = normalize_layout(raw)
        spans = []
        heading_stack: List[Tuple[int, str]] = []
        offset = 0
        for i, paragraph in enumerate(re.split(r"\n\s*\n", raw)):
            paragraph = paragraph.strip()
            if not paragraph:
                continue
            heading = RX_HEADING.fullmatch(paragraph)
            if heading:
                level = len(heading.group(1))
                title = heading.group(2).strip()
                heading_stack = [x for x in heading_stack if x[0] < level]
                heading_stack.append((level, title))
                offset += len(paragraph) + 2
                continue
            spans.append(
                DocumentSpan(
                    text=paragraph,
                    source="plain_text",
                    block_index=i,
                    heading_path=" > ".join(x[1] for x in heading_stack)
                    or "Document Root",
                    start_offset=offset,
                    end_offset=offset + len(paragraph),
                )
            )
            offset += len(paragraph) + 2
        return spans

# ============================================================================
# CANDIDATE GENERATION
# ============================================================================
def fuzzy_label_score(token: str) -> Tuple[Optional[str], float]:
    cleaned = re.sub(r"[^A-Za-z]", "", token).casefold()
    if not cleaned:
        return None, 0.0
    best_label = None
    best = 0.0
    for canonical in PROMPT_LABEL_CANONICAL:
        score = float(fuzz.ratio(cleaned, canonical))
        if score > best:
            best = score
            best_label = canonical
    return best_label, best

def prompt_features(content: str, detector_hits: Sequence[str]) -> Dict[str, float]:
    words = re.findall(r"\b\w+\b", content)
    word_count = len(words)
    imperative_hits = len(RX_IMPERATIVE.findall(content))
    constraint_hits = len(RX_CONSTRAINT.findall(content))
    output_hits = len(RX_OUTPUT.findall(content))
    role_hits = len(RX_ROLE.findall(content))
    input_hits = len(RX_INPUT.findall(content))
    sequence_hits = len(RX_SEQUENCE.findall(content))
    placeholders = len(RX_PLACEHOLDER.findall(content))
    lines = [ln.strip() for ln in content.splitlines() if ln.strip()]
    list_lines = sum(bool(re.match(r"^[-*+]|\d+[.)]", ln)) for ln in lines)
    return {
        "word_count": float(word_count),
        "imperative_hits": float(imperative_hits),
        "constraint_hits": float(constraint_hits),
        "output_hits": float(output_hits),
        "role_hits": float(role_hits),
        "input_hits": float(input_hits),
        "sequence_hits": float(sequence_hits),
        "placeholder_hits": float(placeholders),
        "list_ratio": list_lines / max(len(lines), 1),
        "explicit_detector": float(
            any(
                x
                in {
                    "explicit_label",
                    "fuzzy_prompt_label",
                    "fenced_prompt",
                    "json_prompt_field",
                    "csv_prompt_field",
                }
                for x in detector_hits
            )
        ),
    }

def prompt_probability(features: Dict[str, float]) -> float:
    z = -2.35
    z += 1.65 * features["explicit_detector"]
    z += min(features["imperative_hits"], 4) * 0.38
    z += min(features["constraint_hits"], 4) * 0.30
    z += min(features["output_hits"], 4) * 0.31
    z += min(features["role_hits"], 2) * 0.55
    z += min(features["input_hits"], 3) * 0.23
    z += min(features["sequence_hits"], 3) * 0.22
    z += min(features["placeholder_hits"], 5) * 0.18
    z += min(features["list_ratio"], 1.0) * 0.45
    wc = features["word_count"]
    if 20 <= wc <= 1200:
        z += 0.35
    if 60 <= wc <= 600:
        z += 0.25
    return 1.0 / (1.0 + math.exp(-z))

def utility_score(features: Dict[str, float]) -> int:
    score = 0
    score += min(int(features["imperative_hits"]), 4) * 4
    score += min(int(features["constraint_hits"]), 5) * 5
    score += min(int(features["output_hits"]), 5) * 5
    score += min(int(features["input_hits"]), 4) * 3
    score += min(int(features["role_hits"]), 2) * 4
    score += min(int(features["sequence_hits"]), 5) * 3
    score += min(int(features["placeholder_hits"]), 10) * 2
    wc = int(features["word_count"])
    if 80 <= wc <= 600:
        score += 15
    elif 30 <= wc < 80:
        score += 8
    elif 600 < wc <= 1500:
        score += 7
    elif wc < 10:
        score -= 8
    return max(0, min(score, 100))

class CandidateGenerator:
    def generate(self, spans: List[DocumentSpan]) -> List[Candidate]:
        candidates: List[Candidate] = []
        for idx, span in enumerate(spans):
            candidates.extend(self._from_span(span))
            if idx + 1 < len(spans):
                nxt = spans[idx + 1]
                if (
                    span.page == nxt.page
                    or span.page is None
                    or nxt.page is None
                ):
                    combined = f"{span.text}\n\n{nxt.text}".strip()
                    hits = self._structural_hits(combined)
                    if len(hits) >= 2:
                        candidates.append(
                            self._candidate(
                                combined,
                                "adjacent_structural_blocks",
                                span.heading_path,
                                [span, nxt],
                                hits + ["adjacent_merge"],
                            )
                        )
        return self._resolve_overlaps(candidates)

    def _from_span(self, span: DocumentSpan) -> List[Candidate]:
        text = span.text
        out: List[Candidate] = []
        source_lower = span.source.casefold()
        if source_lower.startswith("json_key:"):
            out.append(
                self._candidate(
                    text,
                    "json_prompt_field",
                    span.heading_path,
                    [span],
                    ["json_prompt_field"],
                )
            )
        if source_lower.startswith("csv:"):
            out.append(
                self._candidate(
                    text,
                    "csv_prompt_field",
                    span.heading_path,
                    [span],
                    ["csv_prompt_field"],
                )
            )
        for match in RX_FENCE.finditer(text):
            body = match.group(2).strip()
            if body:
                out.append(
                    self._candidate(
                        body,
                        "fenced_code",
                        span.heading_path,
                        [span],
                        ["fenced_prompt"] + self._structural_hits(body),
                    )
                )
        for match in RX_BLOCKQUOTE.finditer(text):
            body = re.sub(r"(?m)^\s*>\s?", "", match.group(0)).strip()
            if len(body) > 20:
                out.append(
                    self._candidate(
                        body,
                        "blockquote",
                        span.heading_path,
                        [span],
                        ["blockquote"] + self._structural_hits(body),
                    )
                )
        for match in RX_NUMBERED_INLINE.finditer(text):
            out.append(
                self._candidate(
                    match.group(2),
                    "numbered_inline",
                    span.heading_path,
                    [span],
                    ["numbered_inline"],
                )
            )
        for match in RX_BULLET_INLINE.finditer(text):
            out.append(
                self._candidate(
                    match.group(1),
                    "bullet_inline",
                    span.heading_path,
                    [span],
                    ["bullet_inline"],
                )
            )
        label_matches = list(RX_LABEL.finditer(text))
        if label_matches:
            for i, match in enumerate(label_matches):
                start = match.start()
                end = label_matches[i + 1].start() if i + 1 < len(label_matches) else len(text)
                block = text[start:end].strip()
                if len(block) > 8:
                    out.append(
                        self._candidate(
                            block,
                            "explicit_label",
                            span.heading_path,
                            [span],
                            ["explicit_label", match.group("label").casefold()]
                            + self._structural_hits(block),
                        )
                    )
        first_line = next((x.strip() for x in text.splitlines() if x.strip()), "")
        if first_line:
            first_token = re.split(r"[\s:：.\-]+", first_line)[0]
            label, score = fuzzy_label_score(first_token)
            if label and score >= 72:
                out.append(
                    self._candidate(
                        text,
                        "fuzzy_prompt_label",
                        span.heading_path,
                        [span],
                        [f"fuzzy_label:{label}:{score:.0f}", "fuzzy_prompt_label"]
                        + self._structural_hits(text),
                    )
                )
        structural = self._structural_hits(text)
        if len(structural) >= 2:
            out.append(
                self._candidate(
                    text,
                    "structural_prompt",
                    span.heading_path,
                    [span],
                    structural,
                )
            )
        return out

    def _structural_hits(self, text: str) -> List[str]:
        hits = []
        if RX_IMPERATIVE.search(text):
            hits.append("imperative")
        if RX_CONSTRAINT.search(text):
            hits.append("constraints")
        if RX_OUTPUT.search(text):
            hits.append("output")
        if RX_ROLE.search(text):
            hits.append("role")
        if RX_INPUT.search(text):
            hits.append("input")
        if RX_SEQUENCE.search(text):
            hits.append("sequence")
        if RX_PLACEHOLDER.search(text):
            hits.append("placeholder")
        return hits

    def _candidate(
        self,
        content: str,
        source_type: str,
        heading_path: str,
        spans: Sequence[DocumentSpan],
        detector_hits: Sequence[str],
    ) -> Candidate:
        provenance = [asdict(span) for span in spans]
        detector_hits = list(dict.fromkeys(detector_hits))
        features = prompt_features(content, detector_hits)
        return Candidate(
            content=normalize_layout(content),
            source_type=source_type,
            heading_path=heading_path,
            provenance=provenance,
            detector_hits=detector_hits,
            prompt_probability=round(prompt_probability(features), 4),
            utility_score=utility_score(features),
        )

    def _resolve_overlaps(self, candidates: List[Candidate]) -> List[Candidate]:
        priority = {
            "json_prompt_field": 100,
            "csv_prompt_field": 100,
            "explicit_label": 90,
            "fuzzy_prompt_label": 85,
            "fenced_code": 80,
            "structural_prompt": 70,
            "adjacent_structural_blocks": 60,
            "blockquote": 50,
            "numbered_inline": 40,
            "bullet_inline": 40,
        }
        grouped: Dict[str, Candidate] = {}
        for candidate in candidates:
            key = dedup_canonical(candidate.content)
            existing = grouped.get(key)
            if existing is None:
                grouped[key] = candidate
                continue
            existing.detector_hits = list(
                dict.fromkeys(existing.detector_hits + candidate.detector_hits)
            )
            existing.provenance.extend(candidate.provenance)
            if priority.get(candidate.source_type, 0) > priority.get(existing.source_type, 0):
                existing.source_type = candidate.source_type
            existing.prompt_probability = max(
                existing.prompt_probability,
                candidate.prompt_probability,
            )
            existing.utility_score = max(
                existing.utility_score,
                candidate.utility_score,
            )
        return list(grouped.values())

# ============================================================================
# FORENSIC RISK CLASSIFICATION
# ============================================================================
def forensic_views(content: str) -> Dict[str, str]:
    unicode_view = normalize_unicode(content)
    whitespace_view = re.sub(r"\s+", " ", unicode_view)
    deobfuscated = whitespace_view
    deobfuscated = re.sub(r"[\u200b-\u200f\u2060\ufeff]", "", deobfuscated)
    deobfuscated = deobfuscated.replace("／", "/").replace("＼", "\\")
    return {
        "raw": content,
        "unicode_normalized": unicode_view,
        "whitespace_normalized": whitespace_view,
        "deobfuscated_lexical": deobfuscated,
    }

def classify_risk(content: str) -> Tuple[str, List[str], List[Dict[str, Any]]]:
    views = forensic_views(content)
    evidence = []
    flags = set()
    for view_name, value in views.items():
        for category, patterns in COMPILED_RISK_RULES.items():
            for pattern in patterns:
                for match in pattern.finditer(value):
                    flags.add(category)
                    evidence.append(
                        {
                            "category": category,
                            "view": view_name,
                            "matched_text": match.group(0),
                            "start": match.start(),
                            "end": match.end(),
                        }
                    )
    severe = {
        "command_execution",
        "external_exfiltration",
        "credential_targeting",
        "privilege_escalation",
    }
    count = len(flags)
    severe_count = len(flags.intersection(severe))
    if severe_count >= 2 or count >= 5:
        level = "high"
    elif severe_count >= 1 or count >= 2:
        level = "medium"
    elif count == 1:
        level = "low"
    else:
        level = "none"
    return level, sorted(flags), evidence

# ============================================================================
# DEDUPLICATION
# ============================================================================
class Deduplicator:
    def __init__(
        self,
        lexical_threshold: float = 94.0,
        semantic_threshold: float = 0.965,
        model_name: str = "all-MiniLM-L6-v2",
    ):
        self.lexical_threshold = lexical_threshold
        self.semantic_threshold = semantic_threshold
        self.model_name = model_name
        self._model = None

    def _model_or_none(self):
        if SentenceTransformer is None:
            return None
        if self._model is None:
            self._model = SentenceTransformer(self.model_name)
        return self._model

    def run(self, candidates: List[Candidate]) -> List[Candidate]:
        exact: Dict[str, Candidate] = {}
        for candidate in candidates:
            canonical = dedup_canonical(candidate.content)
            h = hashlib.sha256(canonical.encode()).hexdigest()[:16]
            candidate.content_hash = h
            if h not in exact:
                exact[h] = candidate
            else:
                winner = exact[h]
                winner.duplicate_count += 1
                winner.detector_hits = list(
                    dict.fromkeys(winner.detector_hits + candidate.detector_hits)
                )
                winner.provenance.extend(candidate.provenance)
                winner.prompt_probability = max(
                    winner.prompt_probability,
                    candidate.prompt_probability,
                )
                winner.utility_score = max(
                    winner.utility_score,
                    candidate.utility_score,
                )
        representatives = sorted(
            exact.values(),
            key=lambda x: len(x.content),
            reverse=True,
        )
        near_text: List[Candidate] = []
        for candidate in representatives:
            merged = False
            canon = dedup_canonical(candidate.content)
            for rep in near_text:
                rep_canon = dedup_canonical(rep.content)
                length_ratio = min(len(canon), len(rep_canon)) / max(
                    len(canon),
                    len(rep_canon),
                    1,
                )
                if length_ratio < 0.70:
                    continue
                similarity = fuzz.ratio(canon, rep_canon)
                token_similarity = fuzz.token_set_ratio(canon, rep_canon)
                if max(similarity, token_similarity) >= self.lexical_threshold:
                    rep.duplicate_count += candidate.duplicate_count
                    rep.duplicate_ids.append(candidate.content_hash)
                    rep.provenance.extend(candidate.provenance)
                    rep.detector_hits = list(
                        dict.fromkeys(rep.detector_hits + candidate.detector_hits)
                    )
                    rep.prompt_probability = max(
                        rep.prompt_probability,
                        candidate.prompt_probability,
                    )
                    rep.utility_score = max(
                        rep.utility_score,
                        candidate.utility_score,
                    )
                    merged = True
                    break
            if not merged:
                near_text.append(candidate)
        if len(near_text) < 2:
            return near_text
        model = self._model_or_none()
        if model is None:
            return near_text
        texts = [
            f"Heading: {c.heading_path}\nType: {c.source_type}\nPrompt:\n{c.content}"
            for c in near_text
        ]
        embeddings = model.encode(
            texts,
            normalize_embeddings=True,
            convert_to_numpy=True,
            show_progress_bar=False,
        ).astype(np.float32)
        consumed = set()
        final: List[Candidate] = []
        for i, candidate in enumerate(near_text):
            if i in consumed:
                continue
            final.append(candidate)
            for j in range(i + 1, len(near_text)):
                if j in consumed:
                    continue
                other = near_text[j]
                length_ratio = min(len(candidate.content), len(other.content)) / max(
                    len(candidate.content),
                    len(other.content),
                    1,
                )
                if length_ratio < 0.60:
                    continue
                sim = float(np.dot(embeddings[i], embeddings[j]))
                if sim >= self.semantic_threshold:
                    candidate.duplicate_count += other.duplicate_count
                    candidate.duplicate_ids.append(other.content_hash)
                    candidate.provenance.extend(other.provenance)
                    candidate.detector_hits = list(
                        dict.fromkeys(
                            candidate.detector_hits + other.detector_hits
                        )
                    )
                    consumed.add(j)
        return final

# ============================================================================
# ENGINE
# ============================================================================
class PromptRipper:
    def __init__(
        self,
        minimum_prompt_probability: float = 0.35,
        minimum_utility: int = 0,
        keep_low_probability: bool = True,
        dedup_lexical_threshold: float = 94.0,
        dedup_semantic_threshold: float = 0.965,
    ):
        self.minimum_prompt_probability = minimum_prompt_probability
        self.minimum_utility = minimum_utility
        self.keep_low_probability = keep_low_probability
        self.extractor = DocumentExtractor()
        self.generator = CandidateGenerator()
        self.deduplicator = Deduplicator(
            lexical_threshold=dedup_lexical_threshold,
            semantic_threshold=dedup_semantic_threshold,
        )

    def process(self, path: Path) -> Dict[str, Any]:
        spans = self.extractor.extract(path)
        candidates = self.generator.generate(spans)
        for candidate in candidates:
            level, flags, evidence = classify_risk(candidate.content)
            candidate.red_team_risk = level
            candidate.risk_flags = flags
            candidate.risk_evidence = evidence
        candidates = self.deduplicator.run(candidates)
        accepted = []
        low_probability = []
        for candidate in candidates:
            if candidate.utility_score < self.minimum_utility:
                continue
            if candidate.prompt_probability >= self.minimum_prompt_probability:
                accepted.append(candidate)
            elif self.keep_low_probability:
                low_probability.append(candidate)
        accepted.sort(
            key=lambda c: (
                -c.prompt_probability,
                -c.utility_score,
                c.heading_path,
            )
        )
        low_probability.sort(
            key=lambda c: (
                -c.prompt_probability,
                -c.utility_score,
            )
        )

        def serialize(c: Candidate, index: int) -> Dict[str, Any]:
            return {
                "id": f"P{index:04d}",
                "heading_path": c.heading_path,
                "source_type": c.source_type,
                "content": c.content,
                "content_hash": c.content_hash,
                "prompt_probability": c.prompt_probability,
                "utility_score": c.utility_score,
                "red_team_risk": c.red_team_risk,
                "risk_flags": c.risk_flags,
                "risk_evidence": c.risk_evidence,
                "detector_hits": c.detector_hits,
                "duplicate_count": c.duplicate_count,
                "duplicate_ids": c.duplicate_ids,
                "provenance": c.provenance,
            }

        accepted_json = [
            serialize(candidate, i)
            for i, candidate in enumerate(accepted, 1)
        ]
        offset = len(accepted_json)
        low_json = [
            serialize(candidate, offset + i)
            for i, candidate in enumerate(low_probability, 1)
        ]
        return {
            "metadata": {
                "source_file": str(path),
                "span_count": len(spans),
                "accepted_prompt_count": len(accepted_json),
                "low_probability_candidate_count": len(low_json),
                "minimum_prompt_probability": self.minimum_prompt_probability,
                "minimum_utility": self.minimum_utility,
                "classification_policy": (
                    "Forensic classification only. "
                    "Detected source content is preserved verbatim."
                ),
            },
            "prompts": accepted_json,
            "low_probability_candidates": low_json,
        }

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("input_file", type=Path)
    parser.add_argument("output_file", type=Path)
    parser.add_argument("--min-prompt-probability", type=float, default=0.35)
    parser.add_argument("--min-utility", type=int, default=0)
    parser.add_argument(
        "--discard-low-probability",
        action="store_true",
        help="Exclude low-probability forensic candidates from the report.",
    )
    args = parser.parse_args()
    if not args.input_file.is_file():
        sys.exit(f"Input does not exist: {args.input_file}")
    engine = PromptRipper(
        minimum_prompt_probability=args.min_prompt_probability,
        minimum_utility=args.min_utility,
        keep_low_probability=not args.discard_low_probability,
    )
    report = engine.process(args.input_file)
    args.output_file.parent.mkdir(parents=True, exist_ok=True)
    args.output_file.write_text(
        json.dumps(report, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )
    print(
        json.dumps(
            {
                "status": "complete",
                "output": str(args.output_file),
                "prompts": report["metadata"]["accepted_prompt_count"],
                "low_probability_candidates": report["metadata"][
                    "low_probability_candidate_count"
                ],
            }
        )
    )

if __name__ == "__main__":
    main()
